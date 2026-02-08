"""
report.py - Report Generation Module
콘솔 리포트 + Word 문서 생성
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_console_report(df: pd.DataFrame) -> None:
    """콘솔에 요약 리포트 출력"""
    print("\n" + "="*80)
    print("📊 뉴스 모니터링 리포트")
    print("="*80)

    # 컬럼 이름 매핑 (하이브리드 시스템 호환)
    sentiment_col = "sentiment_final" if "sentiment_final" in df.columns else "sentiment"
    danger_col = "danger_final" if "danger_final" in df.columns else "risk_level"
    category_col = "issue_category_final" if "issue_category_final" in df.columns else "category"
    reason_col = "danger_final_rationale" if "danger_final_rationale" in df.columns else "reason"

    # Section A: 우리 브랜드 - 부정 기사 (위험도 높은 순)
    print("\n[A] 우리 브랜드 - 부정 기사 (위험도 높은 순)")
    print("-" * 80)
    our_negative = df[(df["group"] == "OUR") & (df[sentiment_col].str.contains("NEGATIVE", case=False, na=False))].copy()

    # 위험도 정렬 (HIGH > MEDIUM > LOW)
    risk_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2, "상": 0, "중": 1, "하": 2, "-": 3, "": 4}
    our_negative["risk_order"] = our_negative[danger_col].map(risk_order).fillna(4)
    our_negative = our_negative.sort_values(["risk_order", "pub_datetime"], ascending=[True, False])
    our_negative = our_negative.head(10)

    if len(our_negative) == 0:
        print("✅ 부정 기사 없음")
    else:
        for idx, row in our_negative.iterrows():
            risk_display = row.get(danger_col, "-")
            category_display = row.get(category_col, "-")
            print(f"\n[{risk_display}] {row['query']} | {category_display}")
            print(f"제목: {row['title']}")
            reason = row.get(reason_col, "")
            if reason and reason not in ["-", ""]:
                print(f"이유: {reason}")
            print(f"날짜: {row['pubDate']}")
            print(f"링크: {row['link']}")

    # Section B: 우리 브랜드 - 긍정 기사
    print("\n" + "-" * 80)
    print("[B] 우리 브랜드 - 긍정 기사 (최신순)")
    print("-" * 80)
    our_positive = df[(df["group"] == "OUR") & (df[sentiment_col].str.contains("POSITIVE", case=False, na=False))]
    our_positive = our_positive.sort_values("pub_datetime", ascending=False).head(10)

    if len(our_positive) == 0:
        print("긍정 기사 없음")
    else:
        for idx, row in our_positive.iterrows():
            category_display = row.get(category_col, "-")
            print(f"\n{row['query']} | {category_display}")
            print(f"제목: {row['title']}")
            print(f"날짜: {row['pubDate']}")
            print(f"링크: {row['link']}")

    # Section C: 경쟁사 하이라이트
    print("\n" + "-" * 80)
    print("[C] 경쟁사 하이라이트 (최신순)")
    print("-" * 80)
    comp = df[df["group"] == "COMPETITOR"]
    comp = comp.sort_values("pub_datetime", ascending=False).head(10)

    if len(comp) == 0:
        print("경쟁사 기사 없음")
    else:
        for idx, row in comp.iterrows():
            sentiment_val = row.get(sentiment_col, "")
            sentiment_icon = ""
            if "POSITIVE" in str(sentiment_val).upper():
                sentiment_icon = "😊"
            elif "NEGATIVE" in str(sentiment_val).upper():
                sentiment_icon = "😟"
            else:
                sentiment_icon = "😐"

            category_display = row.get(category_col, "-")
            risk_display = row.get(danger_col, "")
            risk_info = f" | 위험도: {risk_display}" if risk_display not in ["-", ""] else ""
            print(f"\n{sentiment_icon} {row['query']} | {category_display}{risk_info}")
            print(f"제목: {row['title']}")
            reason = row.get(reason_col, "")
            if reason and reason not in ["-", ""]:
                print(f"이유: {reason}")
            print(f"날짜: {row['pubDate']}")
            print(f"링크: {row['link']}")

    print("\n" + "="*80)


def create_word_report(df: pd.DataFrame, output_path: Path) -> None:
    """
    Word 문서로 리포트 생성
    부정 기사 중 위험도가 높은 것들에 집중
    """
    print(f"\n📄 Word 리포트 생성 중...")

    # 컬럼 이름 매핑 (하이브리드 시스템 호환)
    sentiment_col = "sentiment_final" if "sentiment_final" in df.columns else "sentiment"
    danger_col = "danger_final" if "danger_final" in df.columns else "risk_level"
    category_col = "issue_category_final" if "issue_category_final" in df.columns else "category"
    reason_col = "danger_final_rationale" if "danger_final_rationale" in df.columns else "reason"

    doc = Document()

    # 제목
    title = doc.add_heading('뉴스 모니터링 리포트', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 생성일시
    date_para = doc.add_paragraph()
    date_para.add_run(f'생성일시: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()  # 빈 줄

    # ========== Section 1: 긴급 대응 필요 (위험도 HIGH) ==========
    doc.add_heading('1. 긴급 대응 필요 (위험도: HIGH)', level=1)

    high_risk = df[(df["group"] == "OUR") & (df[danger_col].isin(["HIGH", "상"]))].copy()
    high_risk = high_risk.sort_values("pub_datetime", ascending=False)

    if len(high_risk) == 0:
        doc.add_paragraph("✅ 해당 없음")
    else:
        for idx, row in high_risk.iterrows():
            # 기사 제목
            p = doc.add_paragraph()
            p.add_run(f"🔴 {row['query']}").bold = True
            category_display = row.get(category_col, "-")
            p.add_run(f" | {category_display}")

            # 제목
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("제목: ").bold = True
            p.add_run(row['title'])

            # 이유
            reason = row.get(reason_col, "")
            if reason and reason not in ["-", ""]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("이유: ").bold = True
                run = p.add_run(str(reason))
                run.font.color.rgb = RGBColor(255, 0, 0)

            # 날짜
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("날짜: ").bold = True
            p.add_run(row['pubDate'])

            # 링크
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("링크: ").bold = True
            p.add_run(row['link'])

            doc.add_paragraph()  # 기사 간 간격

    doc.add_page_break()

    # ========== Section 2: 모니터링 필요 (위험도 MEDIUM) ==========
    doc.add_heading('2. 모니터링 필요 (위험도: MEDIUM)', level=1)

    medium_risk = df[(df["group"] == "OUR") & (df[danger_col].isin(["MEDIUM", "중"]))].copy()
    medium_risk = medium_risk.sort_values("pub_datetime", ascending=False).head(15)

    if len(medium_risk) == 0:
        doc.add_paragraph("✅ 해당 없음")
    else:
        for idx, row in medium_risk.iterrows():
            # 기사 제목
            p = doc.add_paragraph()
            p.add_run(f"🟡 {row['query']}").bold = True
            category_display = row.get(category_col, "-")
            p.add_run(f" | {category_display}")

            # 제목
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("제목: ").bold = True
            p.add_run(row['title'])

            # 이유
            reason = row.get(reason_col, "")
            if reason and reason not in ["-", ""]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("이유: ").bold = True
                run = p.add_run(str(reason))
                run.font.color.rgb = RGBColor(255, 165, 0)

            # 날짜
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("날짜: ").bold = True
            p.add_run(row['pubDate'])

            # 링크
            p = doc.add_paragraph(style='List Bullet')
            p.add_run("링크: ").bold = True
            p.add_run(row['link'])

            doc.add_paragraph()  # 기사 간 간격

    doc.add_page_break()

    # ========== Section 3: 경미한 이슈 (위험도 LOW) ==========
    doc.add_heading('3. 경미한 이슈 (위험도: LOW)', level=1)

    low_risk = df[(df["group"] == "OUR") & (df[danger_col].isin(["LOW", "하"]))].copy()
    low_risk = low_risk.sort_values("pub_datetime", ascending=False).head(10)

    if len(low_risk) == 0:
        doc.add_paragraph("✅ 해당 없음")
    else:
        for idx, row in low_risk.iterrows():
            p = doc.add_paragraph(style='List Bullet')
            category_display = row.get(category_col, "-")
            p.add_run(f"{row['query']} | {category_display} | ").bold = True
            p.add_run(row['title'])
            p.add_run(f" ({row['pubDate']})")

    doc.add_page_break()

    # ========== Section 4: 긍정 뉴스 ==========
    doc.add_heading('4. 긍정 뉴스', level=1)

    positive = df[(df["group"] == "OUR") & (df[sentiment_col].str.contains("POSITIVE", case=False, na=False))].copy()
    positive = positive.sort_values("pub_datetime", ascending=False).head(10)

    if len(positive) == 0:
        doc.add_paragraph("긍정 기사 없음")
    else:
        for idx, row in positive.iterrows():
            p = doc.add_paragraph()
            p.add_run(f"😊 {row['query']}").bold = True
            category_display = row.get(category_col, "-")
            p.add_run(f" | {category_display}")

            p = doc.add_paragraph(style='List Bullet')
            p.add_run("제목: ").bold = True
            p.add_run(row['title'])

            p = doc.add_paragraph(style='List Bullet')
            p.add_run("날짜: ").bold = True
            p.add_run(row['pubDate'])

            p = doc.add_paragraph(style='List Bullet')
            p.add_run("링크: ").bold = True
            p.add_run(row['link'])

            doc.add_paragraph()

    doc.add_page_break()

    # ========== Section 5: 경쟁사 동향 ==========
    doc.add_heading('5. 경쟁사 동향', level=1)

    competitor = df[df["group"] == "COMPETITOR"].copy()
    competitor = competitor.sort_values("pub_datetime", ascending=False).head(15)

    if len(competitor) == 0:
        doc.add_paragraph("경쟁사 기사 없음")
    else:
        for idx, row in competitor.iterrows():
            sentiment_val = row.get(sentiment_col, "")
            sentiment_icon = ""
            if "POSITIVE" in str(sentiment_val).upper():
                sentiment_icon = "😊"
            elif "NEGATIVE" in str(sentiment_val).upper():
                sentiment_icon = "😟"
            else:
                sentiment_icon = "😐"

            p = doc.add_paragraph()
            p.add_run(f"{sentiment_icon} {row['query']}").bold = True
            category_display = row.get(category_col, "-")
            p.add_run(f" | {category_display}")

            risk_display = row.get(danger_col, "")
            if risk_display and risk_display not in ["-", ""]:
                p.add_run(f" | 위험도: {risk_display}")

            p = doc.add_paragraph(style='List Bullet')
            p.add_run("제목: ").bold = True
            p.add_run(row['title'])

            reason = row.get(reason_col, "")
            if reason and reason not in ["-", ""]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run("이유: ").bold = True
                p.add_run(str(reason))

            p = doc.add_paragraph(style='List Bullet')
            p.add_run("날짜: ").bold = True
            p.add_run(row['pubDate'])

            doc.add_paragraph()

    # 저장
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Word 리포트 저장: {output_path}")
